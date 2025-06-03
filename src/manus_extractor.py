from docx import Document
import subprocess
import os
import json

def convert_doc_to_docx(input_path, output_dir=None):
    """
    Converts a .doc file to .docx format using LibreOffice.
    """
    if not os.path.isfile(input_path):
        raise FileNotFoundError(f"{input_path} does not exist.")
    
    if output_dir is None:
        output_dir = os.path.dirname(input_path)

    # Run the LibreOffice command-line tool to convert the DOC file to DOCX
    subprocess.run([
        "soffice",
        "--headless",
        "--convert-to", "docx",
        "--outdir", output_dir,
        input_path
    ], check=True)

    file_name = os.path.basename(input_path)
    converted_file_name = os.path.splitext(file_name)[0] + '.docx'
    print(f"CONVERTED: {file_name} -> {converted_file_name}")

    converted_file_path = os.path.join(output_dir, converted_file_name)

    return converted_file_path

def extract_manus_columns(docx_path, section_col_index, manuscript_col_index):
    if not os.path.isfile(docx_path):
        raise FileNotFoundError(f"{docx_path} does not exist.")
    
    # Convert DOC to DOCX if necessary
    if docx_path.lower().endswith('.doc'):
        output_dir = os.path.dirname(docx_path)
        docx_path = convert_doc_to_docx(docx_path, output_dir)

    file_name = os.path.basename(docx_path)

    doc = Document(docx_path)

    if not doc.tables:
        raise ValueError(f"No tables found in {docx_path}.")

    longest_table = max(doc.tables, key=lambda table: len(table.rows))

    extracted_data = []
    for row in longest_table.rows[1:]: # Skip header row
        cells = row.cells
        if section_col_index < len(cells) and manuscript_col_index < len(cells):
            section_text = cells[section_col_index].text.strip()
            
            # Check for bold text in runs or paragraph styles
            is_title = any(
                run.bold is True
                for paragraph in cells[section_col_index].paragraphs
                for run in paragraph.runs if run.bold is not None
            ) or any(
                paragraph.style.font.bold is True
                for paragraph in cells[section_col_index].paragraphs
            )

            manuscript_text = cells[manuscript_col_index].text.strip()
            
            extracted_data.append((section_text, manuscript_text, is_title))
    
    return file_name, extracted_data

def save_extracted_data_to_json(file_name, extracted_data, output_dir): 
    base_name = os.path.splitext(file_name)[0]
    json_file_name = f"{base_name}_manuscript.json"
    output_path = os.path.join(output_dir, json_file_name)

    data_to_save = []
    for section, manuscript, is_title in extracted_data:
        if is_title:
            data_to_save.append({"Hovedoverskrift": section, "Tekst": manuscript})
        else:
            data_to_save.append({"Underoverskrift": section, "Tekst": manuscript})

    # Write the data to a JSON file
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(data_to_save, f, ensure_ascii=False, indent=4)

    print(f"Extracted content saved to {output_path}")
    return output_path

def save_manuscript_to_txt(file_name, extracted_data, output_dir):
    # Ensure the output directory exists
    os.makedirs(output_dir, exist_ok=True)

    # Construct the output file path
    base_name = os.path.splitext(file_name)[0]
    text_file_name = f"{base_name}_manuscript.txt"
    output_path = os.path.join(output_dir, text_file_name)

    # Write the extracted data to the text file
    with open(output_path, "w", encoding="utf-8") as f:
        for section, manuscript, is_title in extracted_data:
            # Determine the title status
            title_status = "Hovedoverskrift" if is_title else "Underoverskrift"
            f.write(f"[{title_status}: {section}]\n")

            # Only write "Tekst" if it is not empty
            if manuscript.strip():
                f.write(f"{manuscript}\n")

            f.write("\n")  # Add a blank line between entries

    print(f"Extracted content saved as text to {output_path}")
    return output_path

def run_manus_extractor(input_path, section_col_idx=1, manuscript_col_idx=3):
    try:
        file_name, manuscript = extract_manus_columns(input_path, section_col_index=section_col_idx, manuscript_col_index=manuscript_col_idx)
        #manus = [{"title": el[0], "content": el[1]} for el in manuscript]
        manus = [el[0] + el[1] for el in manuscript]
        #manus_path = save_manuscript_to_txt(file_name, manuscript, output_dir)
    except Exception as e:
        print(f"Faens error: {e}")

    return manus

def data_to_json(data: dict, path: str):
    if path == None:
        with open("result.json", "w", encoding="utf8") as f:
            json.dump(data, f, ensure_ascii=False)
    else:
        with open(path, "w", encoding="utf8") as f:
            for i in data:
                text = {"text": i}
                json.dump(text, f, ensure_ascii=False)
                f.write("\n")


if __name__ == "__main__":
    #test files
    sys77_doc = "/Users/fadulelwalid/Master_Project/data_clean_test/System 77/Manus/F31R0001-O-OT-0019 - Rev 02 - Id (1278233).docx"
    
    save_dir = "/Users/fadulelwalid/Master_Project/data_clean_test" 

    run_manus_extractor(sys77_doc, save_dir)