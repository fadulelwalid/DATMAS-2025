from llmware.prompts import Prompt, Sources
from llmware.models import PyTorchLoader, HFGenerativeModel
from library import LibraryManager


from llmware.models import PromptCatalog, ModelCatalog
from llmware.configs import LLMWareConfig, PostgresConfig
from llmware.retrieval import Query
import json
import time
from transformers import AutoTokenizer

from docx import Document
import os

class Prompter:
    def __init__(self, config: dict, library_manager: LibraryManager):
        self.prompter: Prompt|None = None # Prompter object handles all llmware model interactions
        self.prompt_list = self.add_custom_prompt() # Updates prompt_list with custom prompts

        self.lib_manager = library_manager # LibraryManager object handles all library interactions

        self.config = config # Should be specific model's config from config.json
        self.cur_modelname: str|None = self.config["config"]["model_name"]

        self.cur_model_path: str = self.config["config"]["model_repo_root"]
        self.cur_tokenizer_path: str = self.config["config"]["tokenizer"]
        self.cur_prompt = self.prompt_list[0] if len(self.prompt_list) > 0 else None
        self.cur_lib: str|None = None

        #Flags
        self.is_loaded = False
        self.text_only = True

        #self.prepare_prompter()

    def update_model(self, model_name: str, **kwargs):
        self.cur_modelname = model_name
        pass

    def set_prompt(self, prompt_name: str):
        # Sets the prompt to be used
        if prompt_name in self.prompt_list:
            self.cur_prompt = prompt_name
            return True

    def get_save_name(self):
        return f"output/{time.time()}_{self.cur_modelname}_{self.cur_lib}_{self.cur_prompt}.json"

    def prepare_prompter(self, model_name: str = None, **kwargs):
        #prompter = Prompt(tokenizer=tokenizer)
        #Prompter(prompter).load_model(model_name, **kwargs)
        if model_name != None:
            self.prompter = Prompt(tokenizer=self.config["local"]["base"]).load_model(model_name, **kwargs)
        else:
            self.prompter = Prompt(tokenizer=self.config["config"]["tokenizer"])
            self.load_model(**kwargs)
        self.is_loaded = True
        return True

    def prepare_lib(self, lib_name: str, text_only: bool = True):
        # Loads the library and filters it by content type
        q = Query(self.lib_manager.load_lib(lib_name)).get_whole_library()
        filtered_q = []
        for i in q:
            if i["content_type"] == "text":
                filtered_q.append(i)
        return filtered_q
    
    def _package_response(self, response: dict, save_path: str|None = None, save: bool = True):
        # Packages the response and saves it to a file
        if save:
            save_name = self.get_save_name() if save_path == None else save_path
            with open(save_name, "w") as f:
                json.dump(response, f)
            convrt_json_to_text(save_name)
        return response
    
    def package_query_from_lib(self, query: list):
        if self.text_only:
            filtered_q = []
            for i in query:
                if i["content_type"] == "text":
                    filtered_q.append(i)
        else: 
            filtered_q = query
        return self._package_query(filtered_q)
    
    def package_query_from_inference(self, query: list):
        query = [{"text": x["llm_response"]} for x in query]
        return self._package_query(query)
    
    def _package_query(self, query: list):
        from custom_inferencer import CustomInferencer
        promptCard= PromptCatalog().lookup_prompt(self.cur_prompt)

        Sources(self.prompter).package_source(query, aggregate_source=True)
        engineered_prompts = []
        for index, prompt in enumerate(self.prompter.source_materials):
            prompt_engineered = self.build_core_prompt(promptCard, context=prompt["text"])
            prompt_engineered = prompt_engineered["core_prompt"]
            prompt_final = self.wrap_custom(prompt_engineered, self.prompter.llm_model.prompt_wrapper, instruction=promptCard["system_message"])
            engineered_prompts.append(prompt_final)
        inferencer = CustomInferencer(self.prompter)
        inferencer.engineered_prompts = engineered_prompts
        return inferencer

    def run_stage_one(self, lib: str):
        if self.is_loaded == False:
            print("Model is not loaded")
            return
        self.cur_lib = lib
        
        query = self.prepare_lib(lib)
        inferencer = self.package_query_from_lib(query)
        response = inferencer.run(self.cur_prompt)
        self.prompter.clear_source_materials()
        return self._package_response(response), response
    
    def run_stage_two(self, input):
        inferencer = self.package_query_from_inference(input)
        response = inferencer.run(self.cur_prompt)
        self.prompter.clear_source_materials()
        return self._package_response(response)


    def load_model(self, model_name: str = None, **kwargs):
        max_output = kwargs.pop("max_output", 1000)

        if model_name == None:
            model_name = self.cur_modelname
        model_card = ModelCatalog().lookup_model_card(model_name)
        prompt_wrapper = model_card.get("prompt_wrapper")
        pt_loader = PyTorchLoader()
        #Loads the model based on config
        hf_tokenizer = pt_loader.get_tokenizer(self.cur_tokenizer_path)
        custom_hf_model = pt_loader.get_generative_model(self.cur_model_path, **kwargs)

        model = HFGenerativeModel(custom_hf_model, hf_tokenizer, prompt_wrapper=prompt_wrapper,
                    model_card=model_card, context_window=model_card["context_window"], instruction_following=True, max_output=max_output)

        self.prompter.llm_model = model
        self.prompter.tokenizer = hf_tokenizer
        self.prompter.llm_name = model_name
        self.prompter.context_window_size = self.prompter.llm_model.max_input_len
        self.prompter.llm_max_output_len = max_output
        return
    
    def list_prompts(self):
        # Returns a list of all custom available prompts.
        prompts = PromptCatalog().list_prompts()
        return [x for x in prompts if x not in self.prompt_list]

    
    def add_custom_prompt(self):
        from prompt_templates import add_custom_prompt
        prompt_list = add_custom_prompt()
        return prompt_list
    
    def build_core_prompt(self, prompt_card=None, prompt_name=None, separator="\n", query=None, context=None,
                          inference_dict=None):

        """ Builds the core prompt from the prompt_card template. """

        if context is None or context == "": print("No context provided, this is fatal")

        if not prompt_card:
            prompt_card = PromptCatalog().lookup_prompt(prompt_name)

        core_prompt = ""

        if prompt_card:
            for keys in prompt_card["run_order"]:

                if keys == "instruction":
                    # special handler
                    instruction = prompt_card[keys]
                    core_prompt += instruction + separator
                elif keys == "example":
                    # Special handler
                    example = prompt_card[keys]
                    core_prompt += f"<|start_example|>{example}<|end_example|>" + separator
                else:
                    if not keys.startswith("$"):
                        core_prompt += prompt_card[keys] + separator
                    else:
                        if keys == "$query":
                            core_prompt += query + separator
                        if keys == "$context":
                            core_prompt += f"<|start_context|>{context}<|end_context|> + {separator}"



        prompt_dict = {"core_prompt": core_prompt, "prompt_card": prompt_card}

        return prompt_dict
    
    def wrap_custom(self, text, wrapper_type, instruction=None):

        """ Builds wrapper on Prompt based on the selected wrapper_type. """

        prompt_out = ""

        if wrapper_type in PromptCatalog().prompt_wrapper_lookup:

            prompt_template = PromptCatalog().prompt_wrapper_lookup[wrapper_type]

            if "system_start" in prompt_template:

                if prompt_template["system_start"] != "":

                    prompt_out += prompt_template["system_start"]
                    if instruction:
                        prompt_out += instruction
                    else:
                        prompt_out += "Du er en hjelpfull assistent"

                    if "system_stop" in prompt_template:
                        prompt_out += prompt_template["system_stop"]

            if "main_start" in prompt_template:

                prompt_out += prompt_template["main_start"]
                prompt_out += text

                if "main_stop" in prompt_template:
                    prompt_out += prompt_template["main_stop"]

            if "start_llm_response" in prompt_template:
                prompt_out += prompt_template["start_llm_response"]

        else:
            prompt_out = text

        return prompt_out

def instantiate_tokenizer(tokenizer_path: str):
    tokenizer = AutoTokenizer.from_pretrained(tokenizer_path,
            torch_dtype="auto",
            device_map = "cpu",
            low_cpu_mem_usage=True)
    return tokenizer
    


def save_manuscript_to_docx(json_file_path: str):
    with open(json_file_path, "r") as f:
        data = json.load(f)

    manuscript_text = data.get("response_text", "")
    if not manuscript_text:
        print("No response_text found in JSON file.")
        return
    
    doc = Document()
    doc.add_heading("Generert Manuskript", level=1)

    # Split into paragraphs
    paragraphs = [p.strip() for p in manuscript_text.split("\n") if p.strip()]
    for para in paragraphs:
        # Detect subheadings and make them bold -- må kanskje endres avhengig strukturen på generert summary 
        if para.lower().startswith("## "):
            doc.add_heading(para, level=2)
        else:
            doc.add_paragraph(para)

    base_name = os.path.splitext(os.path.basename(json_file_path))[0]
    docx_file_path = os.path.join(os.path.dirname(json_file_path), f"{base_name}.docx")

    doc.save(docx_file_path)
    print(f"Manuscript saved to: {docx_file_path}")

def convrt_json_to_text(json_file):
    with open(json_file, "r") as f:
        data = json.load(f)
    if isinstance(data, list):
        output = ""
        for i, entries in enumerate(data):
            if "llm_response" in entries:
                output += entries["llm_response"] + "\n"
    else:
        output = data
    output = output.split("\n")
    with open(f"{json_file}.txt", "w") as f:
        f.write("\n".join(output))


if __name__ == "__main__":
    #from config import setup_config
    #setup_config("master") # -C
    #setup_config("Master_Project") # -F

    convrt_json_to_text("output/manuscript_from_summary_qwen_system_73_tech_docs_1744069314.7414737.json")

    #main()
    #run_deep()
    #run_llama()
